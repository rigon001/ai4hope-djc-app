import os
import pandas as pd
from llama_index.core import StorageContext, VectorStoreIndex, Document, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from docx import Document as DocxDocument
import fitz  # PyMuPDF
from tqdm import tqdm

# --- STEP 1: Configure local embedding model (BGE) ---
EMBED_MODEL_PATH = os.getenv(
    "EMBED_MODEL_PATH",
    "./embed_model/models--BAAI--bge-base-en-v1.5/snapshots/<snapshot-id>",
)
EMBED_DEVICE = os.getenv("EMBED_DEVICE", "cpu")
INDEX_PERSIST_DIR = os.getenv("INDEX_PERSIST_DIR", "./index")
EDUCATION_BASE = os.getenv("EDUCATION_MATERIAL_BASE", "./education-material/djc-education-material")

Settings.embed_model = HuggingFaceEmbedding(
    model_name=EMBED_MODEL_PATH,
    device=EMBED_DEVICE
)

# --- Pilot site to language mapping ---
SITE_LANGUAGE_MAP = {
    "INTRAS": "es",
    "UCC": "en",
    "UKB": "de",
    "UKCM": "sl",
    "UP": "pt",
    "UoL": "en"
}

# ---- STEP 2: Load documents from pilot site folders ----
print(f"Loading education materials from {EDUCATION_BASE}...")
all_documents = []

if not os.path.exists(EDUCATION_BASE):
    print(f"⚠ Warning: {EDUCATION_BASE} not found. Proceeding with empty list.")
    SITE_LANGUAGE_MAP = {}
else:
    # Only process folders that exist
    SITE_LANGUAGE_MAP = {
        "INTRAS": "es",
        "UCC": "en",
        "UKB": "de",
        "UKCM": "sl",
        "UP": "pt",
        "UoL": "en"
    }

for pilot_site in SITE_LANGUAGE_MAP.keys():
    site_folder = os.path.join(EDUCATION_BASE, pilot_site)
    language = SITE_LANGUAGE_MAP[pilot_site]
    
    if not os.path.exists(site_folder):
        print(f"⚠ Folder not found: {site_folder}, skipping...")
        continue
    
    print(f"\n=== Processing {pilot_site} ({language}) ===")
    
    # Check for parquet file in this pilot site folder
    parquet_path = os.path.join(site_folder, "metadata_fixed.parquet")
    if os.path.exists(parquet_path):
        print(f"  Found parquet file: {parquet_path}")
        try:
            df = pd.read_parquet(parquet_path)
            df = df[
                df["plain_text"].notnull() & 
                (df["plain_text"].str.strip() != "") & 
                df["title"].notnull()
            ]
            for _, row in tqdm(df.iterrows(), total=len(df), desc=f"  {pilot_site} parquet"):
                all_documents.append(
                    Document(
                        text=row["plain_text"],
                        metadata={
                            "title": row.get("title", ""),
                            "url": row.get("url", "N/A"),
                            "warc_date": row.get("warc_date"),
                            "record_id": row.get("record_id"),
                            "doc_type": "web",
                            "language": language,
                            "pilot_site": pilot_site
                        }
                    )
                )
            print(f"  ✓ Loaded {len(df)} web documents from parquet")
        except Exception as e:
            print(f"  ⚠ Failed to process parquet: {e}")
    
    for filename in tqdm(os.listdir(site_folder), desc=f"{pilot_site} files"):
        filepath = os.path.join(site_folder, filename)
        
        # Skip directories and parquet files (already processed)
        if os.path.isdir(filepath) or filename.lower().endswith(".parquet"):
            continue
        
        if filename.lower().endswith(".pdf"):
            try:
                with fitz.open(filepath) as pdf_doc:
                    text = "\n".join(page.get_text() for page in pdf_doc)
                    if text.strip():
                        all_documents.append(
                            Document(
                                text=text,
                                metadata={
                                    "title": filename,
                                    "source": filepath,
                                    "doc_type": "pdf",
                                    "language": language,
                                    "pilot_site": pilot_site
                                }
                            )
                        )
            except Exception as e:
                print(f"Failed to process PDF {filename}: {e}")
        
        elif filename.lower().endswith(".docx"):
            try:
                docx_doc = DocxDocument(filepath)
                text = "\n".join(p.text for p in docx_doc.paragraphs if p.text.strip())
                if text:
                    all_documents.append(
                        Document(
                            text=text,
                            metadata={
                                "title": filename,
                                "source": filepath,
                                "doc_type": "docx",
                                "language": language,
                                "pilot_site": pilot_site
                            }
                        )
                    )
            except Exception as e:
                print(f"Failed to process DOCX {filename}: {e}")

# ---- STEP 3: OPTIONAL - Load global web documents from data/metadata_fixed.parquet ----
# (This is for web documents not associated with any specific pilot site)
global_parquet_path = "data/metadata_fixed.parquet"
if os.path.exists(global_parquet_path):
    print("\n=== Loading global web documents from data/metadata_fixed.parquet ===")
    df = pd.read_parquet(global_parquet_path)
    df = df[
        df["plain_text"].notnull() & 
        (df["plain_text"].str.strip() != "") & 
        df["title"].notnull()
    ]
    
    for _, row in tqdm(df.iterrows(), total=len(df), desc="Global web docs"):
        all_documents.append(
            Document(
                text=row["plain_text"],
                metadata={
                    "title": row.get("title", ""),
                    "url": row.get("url", "N/A"),
                    "warc_date": row.get("warc_date"),
                    "record_id": row.get("record_id"),
                    "doc_type": "web",
                    "language": "en",  # Default or detect from URL
                    "pilot_site": "global_web"
                }
            )
        )
    print(f"✓ Loaded {len(df)} global web documents")
else:
    print(f"⚠ {global_parquet_path} not found, skipping global web documents...")

# --- STEP 4: Clean old index and build fresh ---
import shutil
if os.path.exists(INDEX_PERSIST_DIR):
    print(f"\n⚠ Removing old index from {INDEX_PERSIST_DIR} to build fresh...")
    shutil.rmtree(INDEX_PERSIST_DIR)

print(f"\n=== Building unified index with {len(all_documents)} documents ===")
index = VectorStoreIndex.from_documents(all_documents)

# --- STEP 5: Save to disk ---
print(f"Saving index to {INDEX_PERSIST_DIR} ...")
index.storage_context.persist(persist_dir=INDEX_PERSIST_DIR)
print("✓ Done! Index saved with language and pilot_site metadata.")
